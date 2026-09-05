"""Export reviewed planning Markdown with the bundled document runtime.

Usage: python export-masterplan-docx.py SOURCE.md DESTINATION.docx
Requires python-docx from Codex workspace dependencies. Does not deploy the app.
"""
from pathlib import Path
import re, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def inline(p, text):
    text=re.sub(r'\[([^]]+)\]\(([^)]+)\)',r'\1 (\2)',text)
    for i, part in enumerate(re.split(r'\*\*(.*?)\*\*',text)):
        r=p.add_run(part.replace('`',''))
        r.bold=bool(i%2)

def export(source, destination):
    source_path=Path(source)
    is_master='MASTERFAHRPLAN' in source_path.name.upper()
    document_label='MASTERFAHRPLAN UND STATUS' if is_master else 'FORTSETZUNGSÜBERGABE UND STATUS'
    footer_label='Arbeitsfassung v6' if is_master else 'Aktuelle Fortsetzungsübergabe'
    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.7)
    sec.left_margin=sec.right_margin=Inches(.75)
    sec.header_distance=sec.footer_distance=Inches(.3)
    for name in ['Normal','Title','Heading 1','Heading 2','Heading 3','List Bullet','List Number']:
        st=doc.styles[name]; st.font.name='Aptos'; st.font.color.rgb=RGBColor(0,0,0)
    normal=doc.styles['Normal']; normal.font.size=Pt(11)
    normal.paragraph_format.space_after=Pt(7)
    normal.paragraph_format.line_spacing=1.12
    for name,size in [('Title',27),('Heading 1',17),('Heading 2',13),('Heading 3',11.5)]:
        st=doc.styles[name]; st.font.size=Pt(size)
        st.paragraph_format.space_before=Pt(14 if name!='Title' else 0)
        st.paragraph_format.space_after=Pt(7)
        st.paragraph_format.keep_with_next=True
    h=sec.header.paragraphs[0]; h.text=f'LUVIA   |   {document_label}   |   05.09.2026'
    h.style=doc.styles['Normal']; h.runs[0].font.size=Pt(8)
    f=sec.footer.paragraphs[0]; f.alignment=2
    f.add_run(f'{footer_label}  ·  Seite ').font.size=Pt(8)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); f._p.append(fld)
    lines=[]
    for raw in source_path.read_text(encoding='utf8').splitlines():
        line=raw.strip()
        if re.fullmatch(r'<!--.*-->',line):
            continue
        starts_block=bool(re.match(r'^(?:#|\||- |\d+\. )',line))
        if line and not starts_block and lines and lines[-1] and not lines[-1].startswith(('#','|')):
            lines[-1]+=' '+line
        else:
            lines.append(line)
    i=0
    while i<len(lines):
        line=lines[i].strip()
        if not line: i+=1; continue
        if line.startswith('|'):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith('|'):
                cells=[c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not all(re.fullmatch(r'[-: ]+',c) for c in cells): rows.append(cells)
                i+=1
            table=doc.add_table(rows=0, cols=len(rows[0])); table.autofit=False
            widths={4:[.58,.52,1.65,4.25],3:[1.4,3.25,2.35],2:[1.55,5.45]}[len(rows[0])]
            for col,w in zip(table.columns,widths): col.width=Inches(w)
            for n,row in enumerate(rows):
                cells=table.add_row().cells
                for c,txt,w in zip(cells,row,widths):
                    c.width=Inches(w); p=c.paragraphs[0]; inline(p,txt)
                    p.paragraph_format.space_after=Pt(5); p.paragraph_format.space_before=Pt(5)
                    for run in p.runs: run.font.size=Pt(10); run.bold=(n==0) or run.bold
                    tcPr=c._tc.get_or_add_tcPr()
                    margins=OxmlElement('w:tcMar')
                    for side in ['top','left','bottom','right']:
                        el=OxmlElement('w:'+side); el.set(qn('w:w'),'85'); el.set(qn('w:type'),'dxa'); margins.append(el)
                    tcPr.append(margins)
                    borders=OxmlElement('w:tcBorders')
                    for side in ['top','left','bottom','right']:
                        el=OxmlElement('w:'+side); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'D6D9DC'); borders.append(el)
                    tcPr.append(borders)
                    if n==0:
                        shade=OxmlElement('w:shd'); shade.set(qn('w:fill'),'F0F2F3'); tcPr.append(shade)
                trPr=cells[0]._tc.getparent().get_or_add_trPr()
                if n==0: trPr.append(OxmlElement('w:tblHeader'))
                trPr.append(OxmlElement('w:cantSplit'))
            doc.add_paragraph()
            continue
        if line.startswith('#'):
            count=len(line)-len(line.lstrip('#'))
            title=re.sub(r'[^\w\s]',' ',line.lstrip('#').strip(),flags=re.U)
            title=re.sub(r'\s+',' ',title)
            p=doc.add_paragraph(title,'Title' if count==1 else 'Heading '+str(min(count-1,3)))
            # Main sections flow naturally; package headings stay with their first paragraph.
        elif line.startswith('- '): inline(doc.add_paragraph(style='List Bullet'),line[2:])
        elif re.match(r'^\d+\. ',line): inline(doc.add_paragraph(style='List Number'),re.sub(r'^\d+\. ','',line))
        else: inline(doc.add_paragraph(),line)
        i+=1
    doc.core_properties.title=doc.paragraphs[0].text
    doc.core_properties.subject='Konsolidierter Luvia Arbeitsstand und Fortsetzung'
    doc.core_properties.author='Luvia Projekt'
    # The bundled default template may carry a blue Title paragraph border.
    # Strip inherited/direct paragraph borders; table cell borders remain intact.
    for root in [doc.styles.element,doc.element]:
        for border in list(root.iter(qn('w:pBdr'))):
            border.getparent().remove(border)
    Path(destination).parent.mkdir(parents=True,exist_ok=True)
    doc.save(destination)
    print(str(destination))

if __name__=='__main__': export(sys.argv[1],sys.argv[2])
